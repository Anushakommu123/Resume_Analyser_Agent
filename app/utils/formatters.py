"""Format tuned resume dict to Markdown and DOCX."""
from typing import Dict, Any
from pathlib import Path


def format_resume_markdown(data: Dict[str, Any]) -> str:
    """Convert tuned resume dict to Markdown string."""
    lines = []
    profile = data.get("profile") or {}
    if isinstance(profile, dict):
        if profile.get("name"):
            lines.append(f"# {profile['name']}")
        if profile.get("title"):
            lines.append(f"**{profile['title']}**\n")
        if profile.get("email"):
            lines.append(f"Email: {profile['email']}")
        if profile.get("phone"):
            lines.append(f"Phone: {profile['phone']}")
        if profile.get("summary"):
            lines.append(f"\n## Summary\n{profile['summary']}\n")
    lines.append("")

    experience = data.get("experience") or []
    if experience:
        lines.append("## Experience\n")
        for entry in experience:
            if isinstance(entry, dict):
                title = entry.get("title", "")
                company = entry.get("company", "")
                lines.append(f"### {title} at {company}")
                if entry.get("duration"):
                    lines.append(f"*{entry['duration']}*")
                if entry.get("description"):
                    lines.append(str(entry["description"]))
                lines.append("")
            else:
                lines.append(f"- {entry}\n")
        lines.append("")

    education = data.get("education") or []
    if education:
        lines.append("## Education\n")
        for entry in education:
            if isinstance(entry, dict):
                lines.append(f"- **{entry.get('degree', '')}**, {entry.get('institution', '')}")
            else:
                lines.append(f"- {entry}")
        lines.append("")

    skills = data.get("skills") or []
    if skills:
        lines.append("## Skills\n")
        lines.append(", ".join(str(s) for s in skills))
        lines.append("\n")

    certs = data.get("certifications") or []
    if certs:
        lines.append("## Certifications\n")
        for c in certs:
            lines.append(f"- {c}")
        lines.append("")

    return "\n".join(lines)


def format_resume_docx(data: Dict[str, Any], output_path: str) -> None:
    """Write tuned resume dict to a DOCX file."""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("DOCX export requires: pip install python-docx") from e

    doc = Document()
    md_text = format_resume_markdown(data)
    for line in md_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph(line.strip("*"))
            p.runs[0].bold = True if p.runs else None
        else:
            doc.add_paragraph(line)
    doc.save(output_path)
